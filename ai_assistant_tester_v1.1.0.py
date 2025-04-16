import json
import os
import re
import sys
import cv2
import numpy as np
import pandas as pd
import yaml
from PyQt5.QtWidgets import *
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont
from docx import Document
from PyPDF2 import PdfReader
import markdown
import xml.etree.cElementTree as ET
from openai import OpenAI
from paddleocr import PaddleOCR


class MultiSelectComboBox(QWidget):
    def __init__(self, items):
        super().__init__()
        self.combo_button = None
        self.layout = None
        self.items = items
        self.selected_items = []
        self.init_ui()

    def init_ui(self):
        self.layout = QHBoxLayout(self)
        self.combo_button = QPushButton("选择用例设计方法")
        self.combo_button.clicked.connect(self.show_dialog)
        self.layout.addWidget(self.combo_button)

    def show_dialog(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("选择用例设计方法")
        dialog.setGeometry(100, 100, 300, 300)

        list_widget = QListWidget(dialog)
        list_widget.addItems(self.items)
        list_widget.setSelectionMode(QListWidget.MultiSelection)

        ok_button = QPushButton("确定", dialog)
        ok_button.clicked.connect(lambda: self.get_selected_items(list_widget, dialog))

        layout = QVBoxLayout(dialog)
        layout.addWidget(list_widget)
        layout.addWidget(ok_button)

        dialog.exec_()

    def get_selected_items(self, list_widget, dialog):
        self.selected_items = [item.text() for item in list_widget.selectedItems()]
        self.combo_button.setText(", ".join(self.selected_items))
        dialog.accept()

    def get_selected_items_text(self):
        return self.combo_button.text()


class GenerateThread(QThread):
    """ 
    异步生成线程类
    
    该类负责在后台线程中调用大语言模型API生成测试用例，避免阻塞主界面。
    通过信号机制与主界面通信，包括完成信号、错误信号和进度信号。
    
    属性:
        finished (pyqtSignal): 生成完成时发出的信号，携带生成结果
        error (pyqtSignal): 发生错误时发出的信号，携带错误信息
        progress (pyqtSignal): 进度更新信号，携带进度信息
    """
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    progress = pyqtSignal(str)  # 添加进度信号

    def __init__(self, prompt, context, job_area, func_type, design_method):
        """
        初始化生成线程
        
        参数:
            prompt (str): 提示词，指导模型生成测试用例
            context (str): 上下文内容，通常是需求文档的内容
            job_area (str): 行业领域，如"互联网/电子商务"、"保险"等
            func_type (str): 功能类型，如"功能测试用例"、"接口测试用例"
            design_method (str): 设计方法，如"等价类划分"、"边界值分析"等
        """
        super().__init__()
        self.prompt = prompt
        self.context = context
        self.job_area = job_area
        self.func_type = func_type
        self.design_method = design_method
        self.api_key = self._load_api_key()
        self.base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1"
        self.model = "deepseek-r1"
    
    def _load_api_key(self):
        """
        从配置文件加载API密钥
        
        首先尝试从api_config.json文件中加载API密钥，如果加载失败或密钥为空，
        则使用默认密钥（仅用于开发阶段）。
        
        返回:
            str: API密钥
        """
        # 尝试从配置文件加载
        try:
            config_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'api_config.json')
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    api_key = config.get("api_key", "")
                    if api_key:
                        return api_key
                    else:
                        print("警告: API密钥为空，将使用默认值")
        except Exception as e:
            print(f"加载API配置失败: {e}")
        
        # 使用默认配置（这不是一个好的做法，应当仅用于开发阶段）
        return 'sk-58335f1c890445ca9306e8a25f1e15c5'  # 仅作为示例

    def generate_cases(self):
        """
        生成测试用例
        
        该方法负责调用大语言模型API生成测试用例，处理流式响应，
        并返回生成的测试用例内容。
        
        处理流程:
        1. 初始化OpenAI客户端
        2. 处理上下文长度限制
        3. 构建用户消息
        4. 发送API请求并处理流式响应
        5. 提取生成的测试用例内容
        
        返回:
            str: 生成的测试用例内容
            
        异常:
            Exception: API调用出错时抛出异常
        """
        # 初始化OpenAI客户端
        client = OpenAI(
            api_key=self.api_key,
            base_url=self.base_url
        )
        
        # 优化上下文处理
        max_context_len = 15000  # 模型能处理的最大上下文长度
        if len(self.context) > max_context_len:
            print(f"上下文过长({len(self.context)}字符)，将截断至{max_context_len}字符")
            self.context = self.context[:max_context_len] + "...[内容已截断]"
            
        reasoning_content = ""  # 定义完整思考过程
        answer_content = ""  # 定义完整回复
        is_answering = False  # 判断是否结束思考过程并开始回复

        try:
            # 构建用户消息
            user_message = {
                'role': 'user', 
                'content': f'所在行业: {self.job_area}；'
                          f'文档内容： {self.context}； '
                          f'生成的用例类型： {self.func_type}； '
                          f'用例设计方法： {self.design_method}； '
                          f'提示词：{self.prompt}'
            }
            
            # 创建聊天完成请求
            completion = client.chat.completions.create(
                model=self.model,
                messages=[user_message],
                stream=True  # 启用流式响应，实时获取生成内容
            )
            
            print("\n" + "=" * 20 + "思考过程" + "=" * 20 + "\n")
            for chunk in completion:
                # 处理空选择的情况（如使用情况统计）
                if not chunk.choices:
                    if hasattr(chunk, 'usage'):
                        print("\nUsage:")
                        print(chunk.usage)
                    continue
                
                # 处理有选择的情况
                delta = chunk.choices[0].delta
                
                # 处理思考过程
                if hasattr(delta, 'reasoning_content') and delta.reasoning_content is not None:
                    print(delta.reasoning_content, end='', flush=True)
                    reasoning_content += delta.reasoning_content
                    self.progress.emit("正在分析需求...")
                    continue
                
                # 处理回复内容
                if hasattr(delta, 'content') and delta.content:
                    # 标记回复开始
                    if not is_answering:
                        print("\n" + "=" * 20 + "完整回复" + "=" * 20 + "\n")
                        is_answering = True
                        self.progress.emit("正在生成用例...")
                    
                    # 输出回复内容
                    print(delta.content, end='', flush=True)
                    answer_content += delta.content
            
            return answer_content
        except Exception as e:
            error_msg = f"API调用出错: {str(e)}"
            print(error_msg)
            raise Exception(error_msg)

    def _extract_json(self, result):
        """
        提取并处理JSON格式的结果
        
        该方法尝试从模型返回的结果中提取JSON格式的内容。
        首先尝试提取```json```块中的内容，如果失败则尝试提取整个文本中的JSON数组。
        
        参数:
            result (str): 原始结果文本
            
        返回:
            str: 提取的JSON字符串或原始结果
        """
        # 尝试提取```json```块中的内容
        json_match = re.search(r"```json(.*)```", result, re.DOTALL)
        if json_match:
            json_str = json_match.group(1).strip()
            # 校验是否有效的JSON格式
            try:
                json.loads(json_str)
                return json_str
            except json.JSONDecodeError:
                print("提取的JSON格式无效，尝试其他方法")
                
        # 如果没有找到或不是有效JSON，尝试在整个文本中查找有效JSON数组
        if result.strip().startswith("[") and result.strip().endswith("]"):
            try:
                json.loads(result.strip())
                return result.strip()
            except json.JSONDecodeError:
                print("文本中的JSON格式无效，返回原始结果")
                
        # 如果都失败了，返回原始结果
        return result

    def run(self):
        """
        线程执行的主要方法
        
        该方法在线程启动时自动调用，负责生成测试用例并发送完成信号。
        如果发生错误，则发送错误信号。
        """
        try:
            # 生成测试用例
            result = self.generate_cases()
            
            # 尝试提取JSON内容并处理
            if 'json' in result.lower() or (result.strip().startswith("[") and result.strip().endswith("]")):
                result = self._extract_json(result)
                
            # 发送完成信号
            self.finished.emit(result)
        except Exception as e:
            # 发送错误信号
            self.error.emit(str(e))


class DeepSeekTool(QMainWindow):
    def __init__(self):
        super().__init__()
        self.thread = None
        self.design_method = None
        self.export_btn = None
        self.export_combo = None
        self.module_input_table = None
        self.module_input = None
        self.result_area = None
        self.prompt_input = None
        self.file_list = None
        self.preview_area = None
        self.refresh_prompt_btn = None
        self.generate_btn = None
        self.btn_clear_all = None
        self.btn_select_all = None
        self.btn_refresh = None
        self.comboBox = None
        self.combo_kb = None
        self.btn_add_kb = None
        self.method_combo = None
        self.func_choice_combo = None
        self.param_choice_combo = None
        self.func_type = None  # 测试用例类型
        self.module_input_pic = None
        self.init_ui()
        self.knowledge_bases = []
        self.current_dir = ""
        self.job_area = None  # 行业　
        self.load_knowledge_bases()  # 新增：初始化时加载历史记录
        self.setStyleSheet(self.load_stylesheet())
        # 模板固有短语列表，可以扩展
        self.template_phrases = [
            "用户行为分析", "功能细项",
            "互联网搜索需求", "基础数据",
            "基表", "基础数据", "样张模板", "XX批处理",
            "批处理要素说明", "批处理执行频率", "计算公式和案例", "非功能性需求",
            "用户量", "运行维护",
            "终端设备",
            "网络需求", "并发量", "业务量", "终端设备",
            "其他"
        ]
        # 正文部分模糊匹配过滤
        self.content_filter_fuzzy = ["【注：",
                                     "不涉及",
                                     "表格：填写页面字段描述、页面事件和页面规则，字段描述初始数据来源于业务需求，可修改。",
                                     "V1.0",
                                     "举例"]  # 包含的正文部分过滤掉，使用in运算符
        # 正文部分精准匹配过滤
        self.content_filter_exact = ["需求规格说明书",
                                     "版本记录",
                                     "细项类别只能选填"业务数据"或"业务功能"；",
                                     "业务数据的细项名称只能以"信息"、"数据"、"配置"或类似名词为后缀；",
                                     "业务功能的细项名称只能以"信息的管理"、"数据的增删改查"、"配置的维护"等有对象的动词为后缀，新增、删除、查询、修改可单列为一个细项；",
                                     "变化说明一般选"新增"、"已有但需修改"，确实没有变化的业务数据或业务功能，可填无变化，也可不填写该细项；",
                                     "备注用做补充说明，按需填写即可；",
                                     "接口功能和批处理功能在相应章节描述即可，不需要在此填写】"]  # 包含的正文部分过滤掉，使用==运算符
        self.clean_flag = ["需求规格说明书", ]

    def init_ui(self):
        """ 初始化界面 """
        self.setWindowTitle("DeepSeek 智能测试分析工具")
        self.setGeometry(300, 200, 1400, 900)

        # 主控件
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout()
        main_widget.setLayout(layout)

        # 创建提示词模式下拉框
        self.param_choice_combo = QComboBox()
        self.param_choice_combo.addItems(["文档", "参数输入"])  # 下拉框选项
        self.param_choice_combo.setCurrentIndex(0)  # 默认选择"文档"

        # 功能模式下拉框
        self.func_choice_combo = QComboBox()
        self.func_choice_combo.addItems(["功能测试用例",
                                         "接口测试用例"])  # 下拉框选项
        self.func_choice_combo.setCurrentIndex(0)  # 默认选择"文档"

        # 用例设计方法多选下拉框
        design_methods = [
            "无",
            "等价类划分",
            "边界值分析",
            "决策表",
            "状态转换",
            "错误推测",
            "场景法",
            "因果图测试",
            "正交分析法"
        ]
        self.method_combo = MultiSelectComboBox(design_methods)  # 测试用例设计方法，支持多选

        industries = [
            "无",
            "互联网/电子商务",
            "保险",
            "金融科技",
            "医疗健康",
            "教育科技",
            "游戏开发",
            "物联网",
            "人工智能",
            "大数据",
            "云计算",
            "汽车电子"
        ]

        # 知识库选择区域
        kb_layout = QHBoxLayout()
        self.btn_add_kb = QPushButton("添加知识库")
        self.combo_kb = QComboBox()
        # 行业
        self.comboBox = QComboBox(self)
        # 将选项添加到 QComboBox
        self.comboBox.addItems(industries)
        self.comboBox.currentTextChanged.connect(self.update_label)
        # 设置默认值为第一个选项
        self.comboBox.setCurrentIndex(0)
        self.btn_refresh = QPushButton("刷新")
        kb_layout.addWidget(QLabel("知识库目录:"))
        kb_layout.addWidget(self.combo_kb)
        kb_layout.addWidget(QLabel("行业:"))
        kb_layout.addWidget(self.comboBox)
        kb_layout.addWidget(QLabel("提示词模式:"))
        kb_layout.addWidget(self.param_choice_combo)
        kb_layout.addWidget(QLabel("功能模式:"))
        kb_layout.addWidget(self.func_choice_combo)
        kb_layout.addWidget(QLabel("用例设计:"))
        kb_layout.addWidget(self.method_combo)

        kb_layout.addWidget(self.btn_add_kb)
        kb_layout.addWidget(self.btn_refresh)

        # 文件操作区域
        file_ops_layout = QHBoxLayout()
        self.btn_select_all = QPushButton("全选")
        # self.btn_clean_docx = QPushButton("清洗")
        self.btn_clear_all = QPushButton("清空")
        file_ops_layout.addWidget(self.btn_select_all)
        # file_ops_layout.addWidget(self.btn_clean_docx)
        file_ops_layout.addWidget(self.btn_clear_all)
        # 生成按钮
        self.generate_btn = QPushButton("开始推理")
        self.generate_btn.setObjectName("generateButton")
        file_ops_layout.addWidget(self.generate_btn)
        # 更新提示词
        self.refresh_prompt_btn = QPushButton("更新提示词")
        self.refresh_prompt_btn.setObjectName("refreshPromptButton")
        file_ops_layout.addWidget(self.refresh_prompt_btn)
        # 文件列表
        self.file_list = QListWidget()
        self.file_list.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.file_list.setSortingEnabled(True)

        # 内容预览
        self.preview_area = QTextEdit()
        self.preview_area.setFixedHeight(200)
        self.preview_area.setReadOnly(False)
        # 提示词输入
        self.prompt_input = QTextEdit()
        self.prompt_input.setFixedHeight(300)
        self.prompt_input.setText("Role: 测试用例设计专家\n\n"
                                  "Rules:\n\n"
                                  "设计目标：\n"
                                  "通过正交分析法实现：\n"
                                  "使用正交表生成参数组合，覆盖所有参数对的交互组合\n\n"
                                  "用例数量：\n"
                                  "尽可能多（不少于15条）\n"
                                  "需求分析指南：\n"
                                  "1. 识别功能边界（系统做什么/不做什么）\n"
                                  "2. 提取业务规则（计算规则、验证规则）\n"
                                  "3. 定义用户角色及其权限\n"
                                  "4. 梳理关键业务流程（正常流、备选流、异常流）\n"
                                  "5. 标记敏感操作（审计日志、权限校验点）\n\n"
                                  "输出要求：\n"
                                  "1. 格式：结构化JSON\n"
                                  "2. 字段：\n"
                                  "   - 用例编号：<模块缩写>-<3位序号>\n"
                                  "   - 用例标题：<测试目标> [正例/反例]\n"
                                  "   - 前置条件：初始化状态描述\n"
                                  "   - 测试数据：参数值的具体组合\n"
                                  "   - 操作步骤：带编号的明确步骤\n"
                                  "   - 预期结果：可验证的断言\n"
                                  "   - 优先级：P0(冒烟)/P1(核心)/P2(次要)\n"
                                  "3. 示例：\n"
                                  "[\n"
                                  "    {\n"
                                  "        \"用例编号\": \"PAY-001\",\n"
                                  "        \"用例标题\": \"支付功能 [正例]\",\n"
                                  "        \"前置条件\": \"用户已登录，购物车内已有商品\",\n"
                                  "        \"测试数据\": {\n"
                                  "            \"支付方式\": \"支付宝支付\",\n"
                                  "            \"金额范围\": \"100-1000\",\n"
                                  "            \"货币类型\": \"CNY\"\n"
                                  "        },\n"
                                  "        \"操作步骤\": [\n"
                                  "            \"1. 打开购物车页面\",\n"
                                  "            \"2. 点击结算按钮\",\n"
                                  "            \"3. 选择支付方式为支付宝支付\",\n"
                                  "            \"4. 确认支付金额为100-1000元人民币\",\n"
                                  "            \"5. 点击支付按钮\"\n"
                                  "        ],\n"
                                  "        \"预期结果\": \"支付成功，页面显示支付完成信息，余额扣减正确\",\n"
                                  "        \"优先级\": \"P1\"\n"
                                  "    }\n"
                                  "]\n\n"
                                  "质量标准：\n"
                                  "- 参数对组合覆盖率 ≥95%\n"
                                  "- 正向场景用例占比60%\n"
                                  "- 异常场景用例占比30%\n"
                                  "- 边界场景用例占比10%\n\n"
                                  "生成步骤：\n"
                                  "1. 参数建模 → 2. 场景分析 → 3. 用例生成 → 4. 交叉校验")

        # 结果展示
        self.result_area = QTextEdit()
        self.result_area.setReadOnly(True)
        self.result_area.setFixedHeight(300)

        self.module_input = QLineEdit()
        self.module_input.setText("需求背景,需求描述,触发")  # 需要指定的标题及其正文，以英文符号分隔
        # 结果展示
        self.result_area = QTextEdit()
        self.result_area.setReadOnly(True)

        # 组装布局
        layout.addLayout(kb_layout)
        hbox = QHBoxLayout()
        hbox.addWidget(QLabel("文本标题:"))
        self.module_input = QLineEdit()
        self.module_input.setText("需求背景,功能描述,触发")
        hbox.addWidget(self.module_input)
        hbox.addWidget(QLabel("表格标题:"))
        self.module_input_table = QLineEdit()
        self.module_input_table.setText("")
        hbox.addWidget(self.module_input_table)
        hbox.addWidget(QLabel("图片标题:"))
        self.module_input_pic = QLineEdit()
        self.module_input_pic.setText("")
        hbox.addWidget(self.module_input_pic)
        layout.addLayout(hbox)
        layout.addWidget(QLabel("文档列表:"))
        layout.addLayout(file_ops_layout)
        layout.addWidget(self.file_list)
        layout.addWidget(QLabel("内容预览:"))
        layout.addWidget(self.preview_area)
        # 底部布局修改
        bottom_layout = QHBoxLayout()
        bottom_layout.addWidget(QLabel("提示词:"))
        bottom_layout.addWidget(self.prompt_input)
        bottom_layout.addWidget(QLabel("生成结果:"))
        bottom_layout.addWidget(self.result_area)

        # 组装主布局
        layout.addLayout(bottom_layout)

        # 导出区域
        export_layout = QHBoxLayout()
        self.export_combo = QComboBox()
        self.export_combo.addItems(["Word 文档", "Text 文件", "Markdown", "JSON", "XLSX"])
        self.export_btn = QPushButton("导出结果")
        export_layout.addWidget(QLabel("导出格式:"))
        export_layout.addWidget(self.export_combo)
        export_layout.addWidget(self.export_btn)

        layout.addLayout(export_layout)

        # 信号连接
        self.btn_add_kb.clicked.connect(self.add_knowledge_base)
        self.btn_refresh.clicked.connect(self.load_knowledge_bases)
        self.combo_kb.currentTextChanged.connect(self.load_directory)
        self.file_list.itemSelectionChanged.connect(self.update_preview)
        self.btn_select_all.clicked.connect(lambda: self.file_list.selectAll())
        # self.btn_clean_docx.clicked.connect(self.clean_text)  # todo
        self.btn_clear_all.clicked.connect(lambda: self.file_list.clearSelection())
        self.generate_btn.clicked.connect(self.generate_report)
        self.refresh_prompt_btn.clicked.connect(self.generate_testcase_prompt)
        self.export_btn.clicked.connect(self.export_result)

    def update_label(self):
        # 更新标签显示当前选中的值
        # old_text = self.prompt_input.toPlainText()
        self.job_area = self.comboBox.currentText()

    def generate_testcase_prompt(self, params=None):
        """
        生成测试用例设计提示词的智能函数

        参数：
        params : dict/list - 参数维度字典或需求文档类型

        返回：
        str - 结构化提示词模板
        """
        # ========== 参数处理模块 ==========
        method = self.method_combo.get_selected_items_text()
        method_list = []  # 已选方法列表
        if method not in ('选择用例设计方法', '无'):
            method_list = method.split(',')
        elif method == '选择用例设计方法' or method == '无':
            method = '常用测试用例设计方法'
        parameters = ""
        func_type = self.func_choice_combo.currentText()
        if func_type == '接口测试用例':
            prompt = """你是一位资深的接口测试工程师，精通各种接口测试方法和技巧。 你的任务是根据提供的接口定义文档，生成全面、有效的接口测试用例。  
请严格按照以下要求执行：  

1.  **理解接口定义**：  
*   仔细阅读并理解提供的接口定义文档，包括接口名称、描述、请求方法、URL、请求参数、请求头、响应状态码、响应体等信息。  
*   识别接口的功能、输入、输出和潜在的业务逻辑。  

2.  **测试用例设计**：  
*   **全面性**： 针对接口的各种场景设计测试用例，包括正常情况、边界情况、异常情况、错误情况等。  
*   **有效性**： 确保每个测试用例都能够有效地验证接口的功能和性能。  
*   **覆盖率**： 尽量覆盖接口的所有功能点和可能的输入组合。  
*   **可读性**： 测试用例描述清晰简洁，易于理解和执行。  

3.  **测试用例类型**： 至少包含以下类型的测试用例：  
*   **功能测试**：  
    *   验证接口是否能够按照预期实现其功能。  
    *   包括正常流程测试、异常流程测试、边界值测试、等价类划分测试等。  
*   **数据验证**：  
    *   验证接口对输入数据的校验是否正确，例如数据类型、格式、范围、长度等。  
    *   包括必填参数校验、可选参数校验、非法参数校验等。  
*   **安全性测试**：  
    *   验证接口是否存在安全漏洞，例如 SQL 注入、XSS 攻击、身份验证绕过等。  
    *   包括权限验证、输入过滤、敏感数据加密等。  
*   **性能测试**：  
    *   验证接口的性能是否满足要求，例如响应时间、吞吐量、并发用户数等。  
    *   包括压力测试、负载测试、稳定性测试等。  
*   **错误处理**：  
    *   验证接口在出现错误时是否能够正确处理，并返回合适的错误信息。  
    *   包括无效参数、服务器错误、网络错误等。  

4.  **输出格式**：  
*   以 JSON 格式输出测试用例，每个测试用例包含以下字段：  
    *   `case_id`: 测试用例的唯一标识符 (例如： TC_001, TC_002...)  
    *   `case_name`: 测试用例的名称 (描述测试目的)  
    *   `priority`: 测试用例的优先级 (High, Medium, Low)  
    *   `pre_condition`: 执行测试用例的前提条件  
    *   `steps`:  测试步骤 (列表，每一步骤包含详细的操作和预期结果)  
    *   `expected_result`: 预期结果  
    *   `test_data`: 测试数据 (JSON 格式)  
    *   `test_type`: 测试类型 (Functional, Data Validation, Security, Performance, Error Handling)  
*   提供至少 3 个测试用例。  

5.  **示例**：  
*   以下是一个测试用例的 JSON 格式示例：  

```json  
{  
"case_id": "TC_001",  
"case_name": "验证用户ID为有效值时，接口返回正确的用户信息",  
"priority": "High",  
"pre_condition": "用户已注册",  
"steps": [  
    "发送 GET 请求到 /users/123",  
    "检查响应状态码是否为 200",  
    "检查响应体是否包含用户ID、用户名、邮箱等信息"  
],  
"expected_result": "接口返回用户ID为123的用户信息，包括用户名、邮箱等",  
"test_data": {  
    "user_id": 123  
},  
"test_type": "Functional"  
                        }  """
            self.prompt_input.clear()
            self.prompt_input.setText(prompt)
        elif func_type == '功能测试用例':
            parameters += f"输出用例类型{func_type}"
            if isinstance(params, dict) and len(params) > 0:
                # 显式参数模式
                parameters = "参数维度：\n" + "\n".join(
                    [f"▸ {k}：{', '.join(v)}" for k, v in params.items()]
                )
            elif isinstance(params, list):
                # 需求文档类型提示
                doc_type = params[0] if params else "通用需求"
                parameters += f"需求文档类型：{doc_type}\n" + \
                              "请提取以下要素：\n" + \
                              "1. 核心业务实体及其属性\n" + \
                              "2. 关键业务流程步骤\n" + \
                              "3. 状态转换规则\n" + \
                              "4. 输入验证规则\n" + \
                              "5. 错误处理策略"
            else:
                # 默认需求分析模式
                parameters += "需求分析指南：\n" + \
                              "1. 识别功能边界（系统做什么/不做什么）\n" + \
                              "2. 提取业务规则（计算规则、验证规则）\n" + \
                              "3. 定义用户角色及其权限\n" + \
                              "4. 梳理关键业务流程（正常流、备选流、异常流）\n" + \
                              "5. 标记敏感操作（审计日志、权限校验点）"

            method_library = {
                "正交分析法": {
                    "desc": "使用正交表生成参数组合，覆盖所有参数对的交互组合",
                    "steps": ["构建正交表", "优化组合数量", "验证两两覆盖"],
                    "coverage": "参数对组合覆盖率 ≥95%"
                },
                "边界值分析": {
                    "desc": "针对数值型参数测试极值：最小值、略高于最小值、正常值、略低于最大值、最大值",
                    "steps": ["识别边界参数", "生成六点值（min-1,min,min+1,norm,max-1,max）", "处理无效类"],
                    "coverage": "边界条件覆盖率100%"
                },
                "等价类划分": {
                    "desc": "将输入划分为有效/无效类，每个类选取代表值测试",
                    "steps": ["定义有效等价类", "定义无效等价类", "生成代表值"],
                    "coverage": "每个等价类至少1个用例"
                },
                "状态转换": {
                    "desc": "基于状态机模型测试合法/非法转换",
                    "steps": ["绘制状态图", "覆盖所有合法转换", "测试非法转换"],
                    "coverage": "状态转换覆盖率100%"
                },
                "决策表": {
                    "desc": "条件组合的全覆盖测试（适合复杂业务规则）",
                    "steps": ["列出所有条件桩", "构建真值表", "合并相似项"],
                    "coverage": "条件组合覆盖率100%"
                },
                "错误推测": {
                    "desc": "基于经验测试易错点：异常输入、中断测试、并发操作",
                    "steps": ["列出历史缺陷", "分析脆弱模块", "设计非常规操作"],
                    "coverage": "补充覆盖边界外的5%"
                },
                "场景法": {
                    "desc": "模拟用户旅程测试端到端流程",
                    "steps": ["识别主成功场景", "定义扩展场景", "组合异常路径"],
                    "coverage": "主流程覆盖率100%"
                },
                "因果图": {
                    "desc": "分析输入条件的逻辑关系生成用例",
                    "steps": ["识别原因和结果", "构建因果图", "生成判定表"],
                    "coverage": "因果逻辑覆盖率100%"
                }
            }
            desc_str = ''
            # ========== 方法选择 ==========
            if len(method_list) >= 1:

                for method in method_list:
                    selected_method = method_library.get(method, method_library["正交分析法"])
                    desc_str += f"""
使用{method}方法设计用例时要符合：{selected_method['desc']}

关键步骤：
 {chr(10).join([f'{i + 1}. {step}' for i, step in enumerate(selected_method['steps'])])}
示例：
 {self.generate_example(method)} \n

质量标准：
 - {selected_method['coverage']}
 - 正向场景用例占比60%
 - 异常场景用例占比30%
 - 边界场景用例占比10%
 \n
                    """
            # ========== 生成提示词 ==========
            prompt = f"""
Role: 测试用例设计专家

Rules:

设计目标：\n
通过{method}实现：\n

用例数量：\n
尽可能多（不少于15条）\n

用例设计需遵循：\n
{desc_str} \n

参数：\n
{parameters} \n

输出要求：
1. 格式：结构化JSON
2. 字段：
   - 用例编号：<模块缩写>-<3位序号>
   - 用例标题：<测试目标> [正例/反例]
   - 前置条件：初始化状态描述
   - 测试数据：参数值的具体组合
   - 操作步骤：带编号的明确步骤
   - 预期结果：可验证的断言
   - 优先级：P0(冒烟)/P1(核心)/P2(次要)

生成步骤：
1. 参数建模 → 2. 场景分析 → 3. 用例生成 → 4. 交叉校验"""
            self.prompt_input.clear()
            self.prompt_input.setText(prompt)
        # return prompt

    @staticmethod
    def generate_example(method):
        """生成方法对应的示例"""
        examples = {
            "正交分析法": """[
        {
            "用例编号": "PAY-001",
            "用例标题": "支付功能 [正例]",
            "前置条件": "用户已登录，购物车内已有商品",
            "测试数据": {
                "支付方式": "支付宝支付",
                "金额范围": "100-1000",
                "货币类型": "CNY"
            },
            "操作步骤": [
                "1. 打开购物车页面",
                "2. 点击结算按钮",
                "3. 选择支付方式为支付宝支付",
                "4. 确认支付金额为100-1000元人民币",
                "5. 点击支付按钮"
            ],
            "预期结果": "支付成功，页面显示支付完成信息，余额扣减正确",
            "优先级": "P1"
        }
    ]""",
            "边界值分析": """[
        {
            "用例编号": "INPUT-001",
            "用例标题": "输入字段长度校验 [边界值测试]",
            "前置条件": "系统显示用户注册页面",
            "测试数据": {
                "用户名": "1个字符",
                "密码": "8个字符",
                "邮箱": "test@example.com"
            },
            "操作步骤": [
                "1. 打开注册页面",
                "2. 在用户名字段输入1个字符",
                "3. 在密码字段输入8个字符",
                "4. 在邮箱字段输入有效邮箱地址",
                "5. 点击提交按钮"
            ],
            "预期结果": "系统提示注册成功",
            "优先级": "P1"
        },
        {
            "用例编号": "INPUT-002",
            "用例标题": "输入字段长度校验 [反例，超长输入]",
            "前置条件": "系统显示用户注册页面",
            "测试数据": {
                "用户名": "超过50个字符",
                "密码": "8个字符",
                "邮箱": "test@example.com"
            },
            "操作步骤": [
                "1. 打开注册页面",
                "2. 在用户名字段输入超过50个字符",
                "3. 在密码字段输入8个字符",
                "4. 在邮箱字段输入有效邮箱地址",
                "5. 点击提交按钮"
            ],
            "预期结果": "系统提示用户名长度超限",
            "优先级": "P2"
        }
    ]""",
            "等价类划分": """[
        {
            "用例编号": "LOGIN-001",
            "用例标题": "登录功能 [有效等价类]",
            "前置条件": "用户已注册",
            "测试数据": {
                "用户名": "valid_user",
                "密码": "correct_password"
            },
            "操作步骤": [
                "1. 打开登录页面",
                "2. 输入用户名为valid_user",
                "3. 输入密码为correct_password",
                "4. 点击登录按钮"
            ],
            "预期结果": "登录成功，跳转到首页",
            "优先级": "P1"
        },
        {
            "用例编号": "LOGIN-002",
            "用例标题": "登录功能 [无效等价类]",
            "前置条件": "用户已注册",
            "测试数据": {
                "用户名": "invalid_user",
                "密码": "random_password"
            },
            "操作步骤": [
                "1. 打开登录页面",
                "2. 输入用户名为invalid_user",
                "3. 输入密码为random_password",
                "4. 点击登录按钮"
            ],
            "预期结果": "登录失败，提示用户名或密码错误",
            "优先级": "P1"
        }
    ]""",
            "状态转换": """[
        {
            "用例编号": "ORDER-001",
            "用例标题": "订单状态转换 [正常流程]",
            "前置条件": "用户购物车内有商品，订单创建成功",
            "测试数据": {
                "初始状态": "已创建",
                "操作": "用户付款"
            },
            "操作步骤": [
                "1. 用户点击付款按钮",
                "2. 系统执行支付操作",
                "3. 支付成功后更新订单状态"
            ],
            "预期结果": "订单状态从'已创建'变为'已支付'",
            "优先级": "P1"
        },
        {
            "用例编号": "ORDER-002",
            "用例标题": "订单状态转换 [非法状态]",
            "前置条件": "订单状态为已取消",
            "测试数据": {
                "初始状态": "已取消",
                "操作": "用户付款"
            },
            "操作步骤": [
                "1. 用户尝试付款已取消的订单",
                "2. 系统拦截付款请求"
            ],
            "预期结果": "操作失败，提示订单已取消，无法付款",
            "优先级": "P2"
        }
    ]""",
            "决策表": """[
        {
            "用例编号": "DISCOUNT-001",
            "用例标题": "会员折扣规则 [决策表测试]",
            "前置条件": "系统具有会员等级和折扣规则",
            "测试数据": {
                "会员等级": "黄金会员",
                "消费金额": "500元"
            },
            "操作步骤": [
                "1. 用户登录账号，确认为黄金会员",
                "2. 添加商品到购物车，消费金额为500元",
                "3. 点击结算按钮"
            ],
            "预期结果": "系统计算折扣，实际支付金额为450元",
            "优先级": "P1"
        }
    ]""",
            "错误推测": """[
        {
            "用例编号": "UPLOAD-001",
            "用例标题": "文件上传 [异常输入]",
            "前置条件": "用户已登录，进入文件上传页面",
            "测试数据": {
                "文件类型": "exe文件",
                "文件大小": "10MB"
            },
            "操作步骤": [
                "1. 用户选择一个exe文件",
                "2. 点击上传按钮"
            ],
            "预期结果": "系统提示不支持的文件类型，上传失败",
            "优先级": "P2"
        }
    ]""",
            "场景法": """[
        {
            "用例编号": "CHECKOUT-001",
            "用例标题": "用户购买商品 [主成功场景]",
            "前置条件": "用户已登录，购物车内有商品",
            "测试数据": {
                "商品": "智能手机",
                "支付方式": "支付宝"
            },
            "操作步骤": [
                "1. 用户进入购物车页面",
                "2. 点击结算按钮",
                "3. 填写收货地址",
                "4. 选择支付方式为支付宝",
                "5. 确认订单并付款"
            ],
            "预期结果": "订单支付成功，显示订单详情",
            "优先级": "P1"
        }
    ]""",
            "因果图": """[
        {
            "用例编号": "LOGIN-003",
            "用例标题": "登录功能 [因果关系测试]",
            "前置条件": "系统有登录模块",
            "测试数据": {
                "用户名": "admin",
                "密码": "correct_password"
            },
            "操作步骤": [
                "1. 用户输入用户名为admin",
                "2. 输入密码为correct_password",
                "3. 点击登录按钮"
            ],
            "预期结果": "登录成功，跳转到管理页面",
            "优先级": "P1"
        },
        {
            "用例编号": "LOGIN-004",
            "用例标题": "登录功能 [因果关系测试 - 异常输入]",
            "前置条件": "系统有登录模块",
            "测试数据": {
                "用户名": "admin",
                "密码": "wrong_password"
            },
            "操作步骤": [
                "1. 用户输入用户名为admin",
                "2. 输入密码为wrong_password",
                "3. 点击登录按钮"
            ],
            "预期结果": "登录失败，提示用户名或密码错误",
            "优先级": "P1"
        }
    ]"""
        }
        return examples.get(method, "此方法示例未实现")

    @staticmethod
    def load_stylesheet():
        """ 加载界面样式表 """
        return """
            QMainWindow {
                background-color: #F0F2F5;
            }
            QComboBox, QLineEdit, QListWidget {
                border: 1px solid #DCDFE6;
                border-radius: 4px;
                padding: 5px;
                min-height: 25px;
            }
            QPushButton {
                background-color: #409EFF;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 15px;
            }
            QPushButton:hover {
                background-color: #66B1FF;
            }
            QPushButton#generateButton {
                background-color: #67C23A;
            }
            QPushButton#generateButton:hover {
                background-color: #85CE61;
            }
            QTextEdit {
                border: 1px solid #DCDFE6;
                border-radius: 4px;
                padding: 10px;
                font-family: Consolas;
            }
        """

    def add_knowledge_base(self):
        """ 添加知识库目录（修改版）"""
        directory = QFileDialog.getExistingDirectory(self, "选择知识库目录")
        if directory:
            if directory not in self.knowledge_bases:
                self.knowledge_bases.append(directory)
                self.combo_kb.addItem(directory)
                self.combo_kb.setCurrentText(directory)
                self.save_knowledge_bases()  # 新增持久化存储

    def load_knowledge_bases(self):
        """ 加载知识库历史记录（修改版）"""
        # 这里应从配置文件读取，示例使用内存存储
        # 示例预设两个测试目录
        self.knowledge_bases = [
            r"C:\Users\Administrator\Documents\知识库"
        ]
        self.combo_kb.addItems(self.knowledge_bases)

    def save_knowledge_bases(self):
        """ 持久化存储知识库目录（示例逻辑）"""
        # 实际应写入配置文件，此处仅示意
        pass

    def load_directory(self, directory):
        """ 加载目录文件（修正版）"""
        try:
            self.current_dir = directory
            self.file_list.clear()

            if not directory:
                return

            if not os.path.exists(directory):
                QMessageBox.warning(self, "路径错误", f"目录不存在: {directory}")
                return

            # 优化文件过滤逻辑
            valid_extensions = ('.docx', '.xlsx', '.md', '.txt', '.pdf', '.json', 'yml', 'yaml')
            for f_name in sorted(os.listdir(directory)):
                full_path = os.path.join(directory, f_name)
                if os.path.isfile(full_path) and f_name.lower().endswith(valid_extensions):
                    item = QListWidgetItem(f_name)
                    item.setData(Qt.UserRole, full_path)
                    item.setToolTip(full_path)  # 添加路径提示
                    self.file_list.addItem(item)

            # 添加数量提示
            self.statusBar().showMessage(f"已加载 {self.file_list.count()} 个文档", 3000)

        except Exception as e:
            QMessageBox.critical(self, "加载错误", str(e))

    @staticmethod
    def clean_headers_footers(content):
        """
        清理页眉、页脚和目录内容
        :param content: 文档内容字典
        :return: 清理后的段落和表格
        """
        cleaned_content = {"paragraphs": [], "tables": content["tables"]}

        # 清理页眉和页脚中的无用内容
        headers_footers = content.get("headers", []) + content.get("footers", [])
        for paragraph in content["paragraphs"]:
            # 如果段落在页眉或页脚中，过滤掉
            if paragraph in headers_footers:
                continue
            # 过滤掉可能是目录的内容
            if re.match(r"^\s*(第[\d一二三四五六七八九十]+章|\d+(\.\d+)*).*$", paragraph):
                continue
            if re.match(r"\\t", paragraph):
                continue
            cleaned_content["paragraphs"].append(paragraph)

        return cleaned_content

    def remove_template_phrases(self, content):
        """
        删除文档中的模板固有内容
        :param content: 文档段落列表
        :return: 清理后的段落列表
        """
        cleaned_paragraphs = []
        for paragraph in content["paragraphs"]:
            if any(phrase in paragraph for phrase in self.template_phrases):
                continue  # 跳过模板内容
            cleaned_paragraphs.append(paragraph)
        return {"paragraphs": cleaned_paragraphs, "tables": content["tables"]}

    def clean_text(self, text):
        """ 文本清洗处理 """
        # 去除标点
        text = self.clean_headers_footers(text)
        return text

    @staticmethod
    def remove_toc(doc):
        """
        去除 Word 文档中的目录
        :param doc:
        """
        paragraphs_to_remove = []

        # 遍历段落，识别目录
        for paragraph in doc.paragraphs:
            if "TOC" in paragraph.style.name:  # 检查样式名称
                paragraphs_to_remove.append(paragraph)
            elif paragraph.text.strip() and (paragraph.text[0].isdigit() or paragraph.text.startswith("1.")):
                # 检查是否为目录条目
                paragraphs_to_remove.append(paragraph)

        # 删除标记的段落
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)
        return doc

    @staticmethod
    def is_needed_title(text, title_list, paragraph):
        for title_keyword in title_list:
            # 匹配标题，判断是否开始捕获
            if title_keyword in text and 'toc' not in paragraph.style.name:
                capture = True
                print("开始捕获正文内容")
                return capture

    # 下面是图片识别部分
    @staticmethod
    def is_heading_enhanced(paragraph):
        heading_features = {
            'keywords': ['标题', 'heading', 'header', 'h1', 'h2', 'h3', 'chapter'],
            'font_size': (14, 72),
            'bold_threshold': 0.7
        }
        # 样式名称检查
        style_match = any(
            kw in (paragraph.style.name or "").lower()
            for kw in heading_features['keywords']
        )
        if style_match:
            return True
        # 格式特征检查
        try:
            font = paragraph.style.font or paragraph.document.styles['Normal'].font
            effective_size = font.size.pt if font.size else 12
            is_bold = font.bold or False
            size_ok = heading_features['font_size'][0] <= effective_size <= heading_features['font_size'][1]
            bold_ok = is_bold >= heading_features['bold_threshold']
            return size_ok and bold_ok
        except Exception as e:
            print(f"格式检查失败: {str(e)}")
            return False

    def get_target_pic(self, file, target_title):
        namespace = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                     'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                     'v': "urn:schemas-microsoft-com:vml",
                     'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
                     'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
                     'pic': "http://schemas.openxmlformats.org/drawingml/2006/picture"}
        # 定义了namespace，建议参看xml.etree.cElementTree中关于namespace的部分，添加了较多，可以自定义修改
        doc = Document(file)
        target_title_list = target_title.strip('').split(',')

        # 打开了对应的word文档，而不是上面讲到的.xml，但是本质是一样的。
        def get_img(root_element, target_tag, target_attribute, out_list):
            # 通过解读word的document.xml，得知与图片相关的rId属性值会出现在两个标签中
            # v:imagedata和a:blip
            # <v:imagedata r:id="rId8" o:title=""/>，<a:blip r:embed="rId8">
            for child in root_element:
                tag = child.tag
                attribute = child.attrib
                if tag in target_tag and target_attribute in child.attrib.keys():
                    target_value = child.attrib[target_attribute]
                    # print(target_value)
                    out_list.append(target_value)
                else:
                    get_img(child, target_tag, target_attribute, out_list)

        xml_element = []
        text_content = []
        found_start = False
        target_title_list = [target_title for target_title in target_title_list if target_title]
        if target_title_list:
            n = 0
            found_times = len(target_title_list)
            for par in doc.paragraphs:
                # doc.paragraphs Proxy object wrapping <w:p> element.
                for title in target_title_list:
                    if title in par.text and 'toc' not in par.style.name.lower() and self.is_heading_enhanced(par):
                        found_start = True
                        n += 1
                        break
                    elif title not in par.text and self.is_heading_enhanced(par):
                        found_start = False
                if found_start and n <= found_times:
                    if par.text == '':
                        print('空白内容')
                    text_content.append(par.text if par.text else '+')
                    xml_element.append(par._element.xml)

        # 此处是表格
        # for tbl in doc.tables:
        #     # print(tbl)
        #     # doc.tables Proxy class for a WordprocessingML <w:tbl> element.
        #     xml_element.append(tbl._element.xml)

        rId = []
        id = []
        for element in xml_element:
            if element:
                root = ET.fromstring(element)
                target_tag = ['{urn:schemas-microsoft-com:vml}imagedata',
                              '{http://schemas.openxmlformats.org/drawingml/2006/main}blip']
                # 即v:imagedata和a:blip
                target_attribute1 = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                target_attribute2 = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'
                # 即:r:embed和r:id
                print('开始获取图片')
                get_img(root, target_tag, target_attribute1, rId)
                get_img(root, target_tag, target_attribute2, id)
        j = 0
        dd = id + rId
        images = []
        for id in dd:
            img_part = doc.part.related_parts[id]
            img_binary = img_part.blob  # 二进制
            img = cv2.imdecode(np.frombuffer(img_binary, np.uint8), cv2.IMREAD_COLOR)  # img_binary 为二进制流图片，此处解码为np数组
            try:
                if not img.any():
                    j += 1
                    continue
                else:
                    img_name = os.path.dirname(file) + r"\img" + str(j) + ".jpg"
                    # 存为jpg
                    cv2.imwrite(img_name, img)
                    images.append(img_name)
                    j += 1
            except ValueError as e:
                print(e)
                img_name = os.path.dirname(file) + r"\img" + str(j) + ".jpg"
                # 存为jpg
                cv2.imwrite(img_name, img)
                images.append(img_name)
                j += 1
        if images:
            return images
        else:
            return None

    @staticmethod
    def perform_ocr_with_paddle(images):
        """
        使用 PaddleOCR 对图片进行文字识别。
        """
        results = []
        # 初始化 PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang="ch")  # 支持方向分类和中英文混合识别
        for image_path in images:
            try:
                img = cv2.imread(image_path)
                ocr_result = ocr.ocr(img, cls=True)
                text_lines = [line[1][0] for line in ocr_result[0]]
                results.append((image_path, "\n".join(text_lines)))
            except Exception as e:
                results.append((image_path, f"OCR 识别失败: {e}"))
        return results

    def extract_text_by_title(self, docx_path, title_keywords, table_keywords, pic_keywords):
        """
        提取多个标题下的正文内容
        :param pic_keywords: 图片路径
        :param table_keywords: 表格路径
        :param docx_path: docx 文件路径
        :param title_keywords: 标题关键词列表
        :return: 提取的正文内容（按标题分组）
        """
        doc = Document(docx_path)
        result = {}
        doc = self.remove_toc(doc)  # 先清理一波目录
        for title_keyword in title_keywords.split(','):
            content = []
            capture = False
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # 判断是否是标题
                    if title_keyword in text and 'toc' not in paragraph.style.name.lower():
                        capture = True
                        content.append(text)
                        continue
                    # 停止捕获正文内容
                    if capture and re.match(r"^\d+(\.\d+)*\s+.+", text):  # 新标题
                        break
                    elif "标题" in paragraph.style.name and capture:  # 遇到下一个标题则停止获取
                        break
                    # 捕获正文内容
                    if capture and text:
                        content.append(text)
            result[title_keyword] = "\n".join(content)
        for pic_keyword in pic_keywords.split(','):
            image_paths = self.get_target_pic(docx_path, pic_keyword)
            print(image_paths)
            if not image_paths:
                print("未找到目标标题下的图片")
            else:
                print(f"提取到 {len(image_paths)} 张图片：{image_paths}")
                # OCR 识别
                print("\n正在进行 OCR 识别...")
                ocr_results = self.perform_ocr_with_paddle(image_paths)
                # 输出结果
                ocr_results_text = ''
                for image_path, text in ocr_results:
                    # print(f"\n图片路径：{image_path}")
                    # print(f"识别内容：\n{text}")
                    ocr_results_text += text + '--'
                result['识别内容'] = ocr_results_text
        return result

    @staticmethod
    def is_title(paragraph, filter_list=None):
        """
        判断段落是否为标题
        :param paragraph: 段落对象
        :param filter_list: 预配置的关键词列表（可选）
        :return: 是否为标题（True/False）
        """
        # 获取段落文本
        text = paragraph.text.strip()

        # 空段落不可能是标题
        if not text:
            return False

        # 样式判断（适用于规范化文档）
        if paragraph.style.name.startswith("Heading"):
            return True

        # 正则表达式匹配标题格式
        if re.match(r"^\d+(\.\d+)*\s+.*", text):  # 数字开头的标题
            return True
        if re.match(r"^\d+-\d+\s+.*", text):  # 带 "-" 的标题
            return True
        if re.match(r"^(附录|参考文献|功能要求|概述).*$", text):  # 特定关键词的标题
            return True

        # 关键词匹配（如果提供了过滤关键词列表）
        if filter_list and any(keyword in text for keyword in filter_list):
            return True

        # 默认不是标题
        return False

    def extract_content(self, file_path, image_folder="extracted_images"):
        """
        提取需求文档内容，包括段落、表格，并处理图片。
        :param file_path: docx 文件路径
        :param image_folder: 图片保存路径（可选）
        :return: dict, 包括段落、表格、页眉和页脚
        """
        doc = Document(file_path)  # 这里可以先清空一波
        doc = self.remove_toc(doc)
        # todo 先把模板里的段落清除一下
        content = {"paragraphs": [], "tables": [], "images": []}
        # 创建图片保存文件夹
        if not os.path.exists(image_folder):
            os.makedirs(image_folder)
        # 提取段落内容
        for paragraph in doc.paragraphs:
            skip_section = False
            text = paragraph.text.strip()

            # 判断是否是标题（假设标题符合 "数字.数字 标题" 格式）
            if self.is_title(paragraph, self.template_phrases):
                # 如果标题中包含过滤列表中的关键词，跳过该段落
                if any(keyword in text for keyword in self.template_phrases) or 'toc' in paragraph.style.name.lower():
                    skip_section = True
                    continue
                else:
                    skip_section = False
            else:  # 此处新增对文本内容的过滤
                for ele in self.content_filter_fuzzy:
                    if ele in text:  # 如果正文过滤内容在正文里，则过滤掉
                        skip_section = True
                for ele in self.content_filter_exact:
                    if ele == text:
                        skip_section = True

                # 如果当前段落不属于需要过滤的章节，则保留
                if not skip_section and text:
                    content["paragraphs"].append(text)

        print(f'----------------\n{content}')

        # 提取表格内容
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            content["tables"].append(table_data)

        # 提取图片 内容（保存到本地）
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_name = os.path.join(image_folder, os.path.basename(rel.target_ref))
                with open(image_name, "wb") as img_file:
                    img_file.write(image_data)
                content["images"].append(image_name)

        return content

    def read_file(self, file_path):
        """ 多格式文件读取 """
        try:
            if file_path.endswith('.docx'):
                if not self.module_input.text():  # ru
                    doc = self.extract_content(file_path)
                else:
                    doc = self.extract_text_by_title(file_path,
                                                     self.module_input.text(),
                                                     self.module_input_table.text(),
                                                     self.module_input_pic.text(),
                                                     )
                return doc
            elif file_path.endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    return '\n'.join([page.extract_text() for page in reader.pages])
            elif file_path.endswith('.xlsx'):
                try:
                    df = pd.read_excel(file_path)
                    # 将 DataFrame 转换为 Markdown 格式
                    excel_content = df.to_markdown(index=False)
                    return excel_content
                except Exception as e:
                    QMessageBox.critical(self, "错误", f"无法读取Excel文件: {str(e)}")
            elif file_path.endswith('.md'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return markdown.markdown(f.read())
            elif file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content
            elif file_path.endswith('.json'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            elif file_path.endswith('.yaml') or file_path.endswith('.yml'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return yaml.safe_load(f)
        except Exception as e:
            QMessageBox.warning(self, "读取错误", f"无法读取文件: {str(e)}")
            return ""

    def update_preview(self):
        """ 更新预览内容 """
        selected = [item.data(Qt.UserRole) for item in self.file_list.selectedItems()]
        content = []
        all_content = ''
        for path in selected:
            raw_text = self.read_file(path)
            if not self.module_input.text() and path.endswith('docx'):
                cleaned = self.clean_text(raw_text)
                content.append(cleaned)
            else:
                content.append({"paragraphs": raw_text})

            if len(content) >= 1 and not self.module_input.text() and path.endswith('docx'):  # 不指定，使用默认配置进行清洗
                for ele in content:
                    all_content += "\n".join(ele["paragraphs"])
            elif len(content) >= 1 and self.module_input.text() and path.endswith('docx'):  # 指定标题获取文档内容
                try:
                    for ele in content:
                        paragraph = ele['paragraphs']
                        for key, value in paragraph.items():
                            all_content += f'{str(value)} \n'  # 获取指定标题内容
                except Exception as e:
                    QMessageBox.critical(self, "预览", f"更新预览内容失败，错误信息{e}！")
            elif len(content) >= 1 and (path.split('.')[-1].lower() in ('txt', 'xlsx')):
                if isinstance(content, list):
                    all_content = content[0].get('paragraphs', '获取内容失败')
            elif len(content) >= 1 and (path.split('.')[-1].lower() in ('json', 'yml', 'yaml')):
                if isinstance(content, list):
                    all_content = content[0].get('paragraphs', '获取内容失败')
                    all_content = json.dumps(all_content, indent=4, ensure_ascii=False)  # 将数据转换为格式化的 JSON 字符串

        self.preview_area.setText(all_content)

    def generate_report(self):
        """ 生成分析报告 """
        if not self.prompt_input.toPlainText().strip():
            QMessageBox.warning(self, "提示", "请输入提示词！")
            return

        context = self.preview_area.toPlainText()
        if not context:
            QMessageBox.warning(self, "提示", "请先选择文档！")
            return

        # 禁用按钮防止重复点击
        self.generate_btn.setEnabled(False)
        QApplication.setOverrideCursor(Qt.WaitCursor)
        self.job_area = self.comboBox.currentText()
        self.func_type = self.func_choice_combo.currentText()
        self.design_method = self.method_combo.get_selected_items_text()

        # 创建异步线程
        self.thread = GenerateThread(
            prompt=self.prompt_input.toPlainText(),
            context=context,
            job_area=self.job_area,
            func_type=self.func_type,
            design_method=self.design_method,
        )
        self.thread.finished.connect(self.on_generation_finished)
        self.thread.error.connect(self.on_generation_error)
        self.thread.start()

    def on_generation_finished(self, result):
        """ 生成完成处理 """
        self.result_area.setText(result)
        self.generate_btn.setEnabled(True)
        QApplication.restoreOverrideCursor()

    def on_generation_error(self, error_msg):
        """ 错误处理 """
        QMessageBox.critical(self, "生成错误", f"模型调用失败:\n{error_msg}")
        self.generate_btn.setEnabled(True)
        QApplication.restoreOverrideCursor()

    @staticmethod
    def json_to_excel(json_data, output_file):
        """
        将任意 JSON 数据中的键作为表头，值作为值，转换为 Excel 表格
        :param json_data: JSON 数据（字符串或字典）
        :param output_file: 输出的 Excel 文件路径
        """
        # 如果输入是 JSON 字符串，将其解析为字典
        data_list = None

        if isinstance(json_data, str):
            json_data = json.loads(json_data)

        # 找到 JSON 数据中的列表部分（假设是字典中的第一个值）
        if isinstance(json_data, dict):
            for key, value in json_data.items():
                if isinstance(value, list):  # 找到第一个值为列表的键
                    data_list = value
                    break
        elif isinstance(json_data, list):
            data_list = json_data
        else:
            raise ValueError("JSON 数据中未找到列表部分！")

        if data_list:
            # 将列表部分转换为 DataFrame
            df = pd.DataFrame(data_list)
            # 将 DataFrame 导出为 Excel 文件
            df.to_excel(output_file, index=False)
            print(f"Excel 文件已成功生成：{output_file}")
        else:
            print("Excel 文件生成失败")

    def export_result(self):
        """ 导出结果 """
        if not self.result_area.toPlainText():
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "保存结果",
            filter=self.get_export_filters()
        )
        if path:
            try:
                content = self.result_area.toPlainText()
                if path.endswith('.docx'):
                    doc = Document()
                    doc.add_paragraph(content)
                    doc.save(path)
                elif path.endswith('.md'):
                    with open(path, 'w') as f:
                        f.write(content)
                elif path.endswith(".json"):
                    with open(path, 'w', encoding='utf-8') as file:
                        file.write(content)
                elif path.endswith(".xlsx"):
                    self.json_to_excel(content, path)
                else:  # txt
                    with open(path, 'w') as f:
                        f.write(content)

                QMessageBox.information(self, "导出成功", "文件已保存！")
            except Exception as e:
                QMessageBox.warning(self, "导出失败", str(e))

    def get_export_filters(self):
        """ 获取导出文件格式过滤器 """
        index = self.export_combo.currentIndex()
        return {
            0: "Word Documents (*.docx)",
            1: "Text Files (*.txt)",
            2: "Markdown Files (*.md)",
            3: "JSON Files (*.json)",
            4: "Excel Files (*.xlsx)"
        }[index]


if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setFont(QFont("微软雅黑", 10))
    window = DeepSeekTool()
    window.show()
    sys.exit(app.exec_())
